"""
responsiva.py — Generador automático de Carta de Responsiva en .docx
Basado en la plantilla: static/responsiva_template.docx
"""

import io
import os
import copy
from datetime import date as _date

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_TEMPLATE = os.path.join(os.path.dirname(__file__), 'static', 'responsiva_template.docx')

_MONTHS_ES = {
    1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
    5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
    9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre',
}

# Asset type display names (maps asset_type → legible label)
_TYPE_LABELS = {
    'laptop':    'Laptop',
    'desktop':   'Desktop / PC',
    'monitor':   'Monitor',
    'tablet':    'Tablet',
    'headset':   'Headset',
    'teclado':   'Teclado',
    'mouse':     'Mouse',
    'impresora': 'Impresora',
    'camara':    'Cámara',
    'otro':      'Otro',
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_text(cell, text, bold=False, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Clear cell content and set text with formatting."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _set_cell_bg(cell, hex_color: str):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_row(table, values: list, bold=False, font_size=10, bg=None, aligns=None):
    """Add a row to the table with given values."""
    row = table.add_row()
    aligns = aligns or [WD_ALIGN_PARAGRAPH.LEFT] * len(row.cells)
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        align = aligns[i] if i < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_text(cell, str(val), bold=bold, font_size=font_size, align=align)
        if bg:
            _set_cell_bg(cell, bg)
    return row


def _clear_table_data_rows(table):
    """Remove all rows except the header (row 0)."""
    for row in list(table.rows[1:]):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)


def _merge_cells_in_row(row, start_col: int, end_col: int):
    """Merge cells from start_col to end_col (inclusive) in a row."""
    row.cells[start_col].merge(row.cells[end_col])


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_responsiva(employee, assets: list, assign_date=None) -> io.BytesIO:
    """
    Generate a Carta de Responsiva Word document.

    Parameters
    ----------
    employee  : Employee model instance
    assets    : list of Asset model instances (the assigned equipment)
    assign_date : datetime.date — defaults to today

    Returns
    -------
    io.BytesIO with the .docx content ready to send_file()
    """
    if assign_date is None:
        assign_date = _date.today()

    doc = Document(_TEMPLATE)

    # ── 1. Date (paragraph index 1) ──────────────────────────────────────────
    date_str = (f"  Torreón, Coahuila, a {assign_date.day} de "
                f"{_MONTHS_ES[assign_date.month]} de {assign_date.year}")
    _replace_paragraph_text(doc.paragraphs[1], date_str)

    # ── 2. Employee name (paragraph index 9) ─────────────────────────────────
    _replace_paragraph_text(doc.paragraphs[9], employee.name.upper())

    # ── 3. Rebuild the equipment table ────────────────────────────────────────
    table = doc.tables[0]
    _clear_table_data_rows(table)

    # Header row already in template (row 0); add COSTO UNITARIO header
    # We need to add a 5th column — we do this by adjusting header
    header_row = table.rows[0]
    # The 4-col header already exists; we'll add cost col text in col 3
    # and add a 5th col below by restructuring

    # Since adding columns to existing Word tables is very complex in python-docx,
    # we keep 4 columns and split: EQUIPO | MARCA | TIPO/MODELO+SERIE | COSTO
    # Re-label header col 3
    _set_cell_text(header_row.cells[3], 'NÚMERO DE SERIE', bold=True,
                   font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # We'll add cost as a 5th column by adding a new cell to each row.
    # Simpler: keep 4 cols, split col 2 for model & serial, add cost at end.
    # BEST approach: recreate the entire table with 5 columns.

    # ── Recreate table as 5 columns ───────────────────────────────────────────
    # We can't easily add columns in python-docx, so we'll delete the existing
    # table and insert a fresh one in the same position.

    _rebuild_table(doc, table, assets, assign_date)

    # ── Save to buffer ─────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _replace_paragraph_text(paragraph, new_text: str):
    """Replace all runs in a paragraph with new_text, preserving first run formatting."""
    # Clear all runs
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def _rebuild_table(doc, old_table, assets: list, assign_date):
    """
    Replace the old 4-column table with a new 5-column table inline.
    Columns: EQUIPO | MARCA | TIPO / MODELO | NÚMERO DE SERIE | COSTO UNITARIO
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Get position of old table in document body
    body = doc.element.body
    old_tbl = old_table._tbl
    old_tbl_idx = list(body).index(old_tbl)

    # Remove old table
    body.remove(old_tbl)

    # Create new table with 5 columns
    new_table = doc.add_table(rows=1, cols=5)
    new_table.style = 'Table Grid'

    # --- Header row ---
    hdr = new_table.rows[0]
    hdr_data = ['EQUIPO', 'MARCA', 'TIPO / MODELO', 'NÚMERO DE SERIE', 'COSTO UNITARIO']
    hdr_aligns = [WD_ALIGN_PARAGRAPH.CENTER] * 5
    for cell, text, align in zip(hdr.cells, hdr_data, hdr_aligns):
        _set_cell_text(cell, text, bold=True, font_size=10, align=align)
        _set_cell_bg(cell, '1F3864')  # dark navy
        # White text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # --- Asset rows ---
    total_cost = 0.0
    for i, asset in enumerate(assets):
        row = new_table.add_row()
        bg = 'F2F4F8' if i % 2 == 0 else 'FFFFFF'

        equipo = _TYPE_LABELS.get(asset.asset_type or 'laptop', asset.asset_type or 'Equipo')
        marca  = (asset.brand.name if asset.brand_id and asset.brand else
                  asset.manufacturer or '—')
        modelo = asset.model or '—'
        serie  = asset.serial_number or '—'
        costo  = asset.purchase_cost or 0.0
        total_cost += costo

        values = [equipo, marca, modelo, serie,
                  f'${costo:,.2f} MXN' if costo else '—']
        aligns = [WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT]
        for cell, val, align in zip(row.cells, values, aligns):
            _set_cell_text(cell, val, font_size=10, align=align)
            _set_cell_bg(cell, bg)

    # --- Total row ---
    total_row = new_table.add_row()
    # Merge first 4 cells
    total_row.cells[0].merge(total_row.cells[3])
    _set_cell_text(total_row.cells[0], 'COSTO TOTAL DEL SETUP',
                   bold=True, font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_bg(total_row.cells[0], 'D9E1F2')

    cost_cell = total_row.cells[1]   # after merge, index shifts
    _set_cell_text(cost_cell, f'${total_cost:,.2f} MXN',
                   bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_bg(cost_cell, '1F3864')
    for p in cost_cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Move new table to original position
    new_tbl = new_table._tbl
    body.remove(new_tbl)
    body.insert(old_tbl_idx, new_tbl)

    # Set column widths
    col_widths_cm = [3.2, 3.0, 3.8, 3.8, 2.8]
    for row in new_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(col_widths_cm[i])
